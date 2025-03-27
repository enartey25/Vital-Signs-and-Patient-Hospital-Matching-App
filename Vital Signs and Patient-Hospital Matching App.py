# These are all the libraries I use for this project. It includes:
# 1. Kivy for the front-end design of the app
# 2. Webbrowser and webdriver from Selenium to open and fill the Google Form and the Google Maps
# 3. Docx for saving vitals in a Word form.
# 4. Geopy for retrieving the position of the end user.

from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.clock import Clock
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.widget import Widget
import webbrowser
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import time
from geopy.geocoders import Nominatim
from math import radians, sin, cos, sqrt, atan2
from kivy.uix.gridlayout import GridLayout
from kivy.lang import Builder
from kivy.uix.image import Image
from time import sleep
import os

Builder.load_file('Vital Signs and Patient-Hospital Matching App.kv')
Builder.load_file('style.kv')


#These are HeFRA (Health Facilities Regulatory Agency) certified hospitals. Below are hospitals I am using to test-run the app.
#the input data is in a 2D list in the format ("name", lat, long, available beds, "Google form link"). These are HeFRA cerified hospitals.
hospitals = [
    ["Korle Bu Teaching Hospital", 5.5381, -0.2272, 2, "Google Form Link 1"],
    ["Komfo Anokye Teaching Hospital", 6.698, -1.629, 2, "Google Form Link 2"],
    ["Tamale Teaching Hospital", 9.393, -0.824, 15, "Google Form Link 3"],
    # More hospitals here if needed
]

#Functions
#All the functions I made specifically for this project.
#The function below is the Haversine formula which calculates the position of the end-user from a place

def haversine_distance(lat1, lon1, lat2, lon2):
    # Convert latitudes and longitudes from degrees to radians for simplicity
    lat1, lon1, lat2, lon2 = map(radians, [lat1, lon1, lat2, lon2])

    # This is the Haversine formula
    dlat = lat2 - lat1
    dlon = lon2 - lon1
    a = sin(dlat / 2) ** 2 + cos(lat1) * cos(lat2) * sin(dlon / 2) ** 2
    c = 2 * atan2(sqrt(a), sqrt(1 - a))
    distance = 6371 * c  # Radius of the Earth in kilometers (Ghana uses the metric system)

    return distance

#The function below gets the coordinates of the end user by using a prompt from the end user
def get_coordinates(location_name):
    geolocator = Nominatim(user_agent="my_app")
    location = geolocator.geocode(location_name)
    if location:
        return location.latitude, location.longitude
    else:
        return None, None


#This function retrieves the vital signs of the patient from a stored Word Document
def recall_vital_signs(patient_name):
    try:
        # Open the saved Word document
        doc = Document(f'{patient_name}_vital_signs.docx')
        
        # Print the vital signs information
        print(f"Vital Signs for {patient_name}:")
        for paragraph in doc.paragraphs:
            print(paragraph.text)
    except FileNotFoundError:
        print(f"No vital signs found for {patient_name}.")

options = Options()
options.add_argument("--headless")
options.add_argument("--disable-gpu")


#Now, the application itself, built using Kivy as stated earlier



class MainMenu(BoxLayout):
    def __init__(self, **kwargs):
        super(MainMenu, self).__init__(**kwargs)
        self.orientation = 'vertical'
        self.spacing = 20
        self.padding = [40, 20, 40, 20]
        self.size_hint = (1, 1)
        self.pos_hint = {'center_x': 0.5, 'center_y': 0.5}
        Clock.schedule_once(self.themainmenu, 1)

    def themainmenu(self, instance):    
        self.clear_widgets()
        
        # Main container
        main_container = BoxLayout(orientation='vertical', spacing=30)
        
        # Header Section
        header = BoxLayout(orientation='vertical', size_hint=(1, 0.3))
        self.image_widget = Image(
            source='assets/Ghana_Health_logo.png',  # Update with actual image path
            allow_stretch=True,
            keep_ratio=True,
            size_hint=(0.99, 2),
            size=(200, 200)
        )
        header.add_widget(self.image_widget)
        
        self.title_label = Label(
            text='[b]Vital Signs App[/b]',
            markup=True,
            font_size='50sp',
            font_name='Roboto',
            color=(0.12, 0.46, 0.70, 1),  # Dark blue
            size_hint=(1, 1),
            height=60
        )
        header.add_widget(self.title_label)
        main_container.add_widget(header)

        # Button Container
        button_container = GridLayout(
            cols=1,
            spacing=25,
            size_hint=(0.8, 0.6),
            pos_hint={'center_x': 0.5}
        )
        
        # Custom Styled Buttons
        menu_items = [
            ('Record Vitals & Find A Hospital', self.collect_vitals),
            ('Recall Vitals', self.recall_vital_signs),
            ('Exit Application', self.exit_app)
        ]

        
        for text, callback in menu_items:
            btn = Button(
                text=text,
                font_size='24sp',
                bold=True,
                background_color=(0.12, 0.46, 0.70, 1),
                background_normal='',
                color=(1, 1, 1, 1),
                size_hint=(1, 1),
                height=80,
                padding=(20, 10)
            )
            btn.bind(on_press=callback)
            button_container.add_widget(btn)

        main_container.add_widget(button_container)
        
        # Footer Section
        footer = Label(
            text='[i]Medical Assistance System[/i]',
            markup=True,
            font_size='20sp',
            color=(0.4, 0.4, 0.4, 1),
            size_hint=(1, 0.1),
            height=5 
        )
        main_container.add_widget(footer)

        footer = Label(
            text='[i]© 2025. Built by Ethan Nartey. Designed by King-Frederick Akyea[/i]',
            markup=True,
            font_size='18sp',
            color=(0.4, 0.4, 0.4, 1),
            size_hint=(1, None),
            height=10 
        )
        main_container.add_widget(footer)

        self.add_widget(main_container)
        


    def collect_vitals(self, instance):
        self.clear_widgets()
        
        layout = BoxLayout(orientation='vertical', spacing=10)  # Add spacing to the layout
        self.add_widget(layout)

        
        # Add some empty space at the top
        self.image_widget = Image(source='assets/Ghana_Health_logo.png', 
                        allow_stretch=True, 
                        keep_ratio=True, 
                        size_hint=(1, 1)
                        )
        layout.add_widget(self.image_widget)
        self.image_widget.pos_hint = {'center_x': 0.5, 'center_y': 0.5}
        
        self.vitals = Label(text="Patient's Details", font_size=64, font_name='Roboto', bold=True, color=(0.12, 0.46, 0.70, 1))
        self.vitals.size_hint = (1, 0.8)
        layout.add_widget(self.vitals)

        self.patient_name_input = TextInput(multiline=False, hint_text='Enter Your Name', font_size=25)
        self.patient_name_input.pos_hint = {'center_x': 0.5, 'center_y': 0.5}  # Center the button
        self.patient_name_input.size_hint = (0.5, 0.5)
        self.patient_name_input.font_size = 25
        self.patient_name_input.height = 100
        layout.add_widget(self.patient_name_input)



        spacer = Widget(size_hint_y=0.01)
        self.add_widget(spacer)

        middle_layout = GridLayout(cols=2, spacing=5)
        layout.add_widget(middle_layout)

        left_middle_layout = BoxLayout(orientation='vertical', spacing=10)
        middle_layout.add_widget(left_middle_layout)

        self.bp_input = TextInput(multiline=False, hint_text='Blood Pressure', font_size=25, size_hint=(1, 0.5))
        left_middle_layout.add_widget(self.bp_input)

        self.temperature_input = TextInput(multiline=False, hint_text='Temperature', font_size=25, size_hint=(1, 0.5))
        left_middle_layout.add_widget(self.temperature_input)

        self.pulse_rate_input = TextInput(multiline=False, hint_text='Pulse Rate', font_size=25, size_hint=(1, 0.5))
        left_middle_layout.add_widget(self.pulse_rate_input)

        right_middle_layout = BoxLayout(orientation='vertical', spacing=10)
        middle_layout.add_widget(right_middle_layout)

        self.oxygen_sat_input = TextInput(multiline=False, hint_text='Oxygen Saturation', font_size=25, size_hint=(1, 0.5))
        right_middle_layout.add_widget(self.oxygen_sat_input)

        self.respiratory_rate_input = TextInput(multiline=False, hint_text='Respiratory Rate', font_size=25, size_hint=(1, 0.5))
        right_middle_layout.add_widget(self.respiratory_rate_input)

        spacer = Widget(size_hint_y=0.09)
        self.add_widget(spacer)      

        self.summary_input = TextInput(multiline=True, hint_text='Additional Information', font_size=30, )
        self.summary_input.pos_hint = {'top': 1, 'left': 0.5, 'center_x': 0.5}
        self.summary_input.size_hint = (0.5, 0.8)
        layout.add_widget(self.summary_input)

        bottom_layout = GridLayout(cols=2, spacing=10)
        layout.add_widget(bottom_layout)

        left_bottom_layout = BoxLayout(orientation='vertical', spacing=10)
        bottom_layout.add_widget(left_bottom_layout)

        self.back_button = Button(text='Main Menu', size_hint=(0.4, None), pos_hint={'center_x': 0.5, 'center_y': 0.5}, background_color=(0.12, 0.46, 0.70, 1))
        self.back_button.bind(on_press=self.themainmenu)
        left_bottom_layout.add_widget(self.back_button)

        right_bottom_layout = BoxLayout(orientation='vertical', spacing=10)
        bottom_layout.add_widget(right_bottom_layout)

        self.submit_button = Button(text='Find a Hospital', size_hint=(0.4, None), pos_hint={'center_x': 0.5, 'center_y': 0.5}, background_color=(0.12, 0.46, 0.70, 1))
        self.submit_button.bind(on_press=self.find_location)
        right_bottom_layout.add_widget(self.submit_button)

        spacer = Widget(size_hint_y=0.09)
        self.add_widget(spacer)

        return layout

    def find_location(self, instance):
        
        self.patient_name = self.patient_name_input.text
        self.bp = self.bp_input.text
        self.temperature = self.temperature_input.text
        self.pulse_rate = self.pulse_rate_input.text
        self.oxygen_sat = self.oxygen_sat_input.text
        self.respiratory_rate = self.respiratory_rate_input.text
        self.summary = self.summary_input.text
        
        
        # Clear the widgets
        self.clear_widgets()

        
        
        # Create a new layout to ask for the person's location
        layout = BoxLayout(orientation='vertical', spacing=30)
        self.add_widget(layout)

        self.image_widget = Image(source='assets/Ghana_Health_logo.png', 
                        allow_stretch=True, 
                        keep_ratio=True, 
                        size_hint=(0.5, 0.5)
                        )
        layout.add_widget(self.image_widget)
        self.image_widget.pos_hint = {'center_x': 0.5, 'center_y': 0.5}

        label = Label(text='Your Location:', font_name='Roboto', bold=True, color=(0.12, 0.46, 0.70, 1), font_size=64)
        label.size_hint = (1, None)  # Make the label smaller
        layout.add_widget(label)

        self.location_input = TextInput(multiline=False, hint_text='Enter Your Current Location', size_hint=(0.3, None), pos_hint={'center_x': 0.5, 'center_y': 0.9}, font_size=25)
        layout.add_widget(self.location_input)

        self.get_hospital_button = Button(text='Find a Hospital', size_hint=(0.3, None), pos_hint={'center_x': 0.5, 'center_y': 0.5}, background_color=(0.12, 0.46, 0.70, 1))
        layout.add_widget(self.get_hospital_button)  # Add the button to the layout
        self.get_hospital_button.bind(on_press=self.display_message)
    
        spacer = Widget(size_hint_y=0.4)
        self.add_widget(spacer)
    
    def display_message(self, instance):
        self.clear_widgets()
        layout = BoxLayout(orientation='vertical', spacing=30)
        self.add_widget(layout)

        self.image_widget = Image(source='assets/Ghana_Health_logo.png', 
                allow_stretch=True, 
                keep_ratio=True, 
                size_hint=(0.5, 0.5)
                )
        layout.add_widget(self.image_widget)
        self.image_widget.pos_hint = {'center_x': 0.5, 'center_y': 0.5}
        label1 = Label(text='''You are being linked to the nearest HeFRA certified hospital. 
Please wait....''', 
                   font_name='Roboto', bold=True, 
                   color=(0, 0, 0, 1), 
                   font_size=35)
        layout.add_widget(label1)
        
        def show_second_message(dt):
            layout.remove_widget(label1)
            label2 = Label(text='''Linked Successfully! 
Your data is being sent to the hospital. Please wait....''', 
                   font_name='Roboto', bold=True, 
                   color=(0, 0, 0, 1), 
                   font_size=35)
            layout.add_widget(label2)
            
            def show_third_message(dt):
                layout.remove_widget(label2)
                label3 = Label(text='''Data Sent Successfully! 
A Google Maps webpage will open to show you the fastest route. Please wait.... 
Get well soon!''', 
                           font_name='Roboto', bold=True, 
                           color=(0, 0, 0, 1), 
                           font_size=35)
                layout.add_widget(label3)
                Clock.schedule_once(self.get_location, 3)
            
            Clock.schedule_once(show_third_message, 3)
        
        Clock.schedule_once(show_second_message, 3)
    
    def get_location(self, instance):
        current_location_name = self.location_input.text
        current_latitude, current_longitude = get_coordinates(current_location_name)
        if current_latitude is None and current_longitude is None:
            self.clear_widgets()
            layout = BoxLayout(orientation='vertical', spacing=30)
            self.add_widget(layout)

            self.image_widget = Image(source='Your Image Here', 
                allow_stretch=True, 
                keep_ratio=True, 
                size_hint=(None, None)
                )
            layout.add_widget(self.image_widget)
            self.image_widget.pos_hint = {'center_x': 0.5, 'center_y': 0.5}
            label = Label(text='''Location not found
You are being redirected to the main menu
Please wait.......''', 
                      font_name='Roboto', bold=True, 
                      color=(0, 0, 0, 1), 
                      font_size=30)
            layout.add_widget(label)
            Clock.schedule_once(self.themainmenu, 3)
        else:
            # Variables for the ideal hospital
            ideal_hospital = None
            ideal_distance = float('inf')

            # Find the ideal hospital
            for hospital in hospitals:
                hospital_name, hospital_latitude, hospital_longitude, available_beds, google_form_link = hospital
                distance = haversine_distance(current_latitude, current_longitude, hospital_latitude, hospital_longitude)
                    

                if available_beds > 0 and distance < ideal_distance:
                    ideal_distance = distance
                    ideal_hospital = hospital_name
                    ideal_google_link = google_form_link
        #Use Google maps for directions to the place
            if ideal_hospital:
                destination_location = f"{ideal_hospital} Hospital"
                url = f"https://www.google.com/maps/dir/{current_latitude},{current_longitude}/{destination_location}"
                webbrowser.open(url)
            else:
                Clock.schedule_once(self.themainmenu, 0.5)

                        #Form gets filled online in the hospital's Google Form Link

            #Initialize web browser

            web = webdriver.Chrome(options=options)
            web.get(ideal_google_link)
            time.sleep(2)

            #Enter details

            enter_name = web.find_element("xpath", '//*[@id="mG61Hd"]/div[2]/div/div[2]/div[1]/div/div/div[2]/div/div[1]/div/div[1]/input')
            enter_name.send_keys(self.patient_name)

            enter_bp = web.find_element("xpath", '//*[@id="mG61Hd"]/div[2]/div/div[2]/div[2]/div/div/div[2]/div/div[1]/div/div[1]/input')
            enter_bp.send_keys(self.bp)

            enter_temp = web.find_element("xpath", '//*[@id="mG61Hd"]/div[2]/div/div[2]/div[3]/div/div/div[2]/div/div[1]/div/div[1]/input')
            enter_temp.send_keys(self.temperature)

            enter_pulse = web.find_element("xpath", '//*[@id="mG61Hd"]/div[2]/div/div[2]/div[4]/div/div/div[2]/div/div[1]/div/div[1]/input')
            enter_pulse.send_keys(self.pulse_rate)

            enter_oxygensat = web.find_element("xpath", '//*[@id="mG61Hd"]/div[2]/div/div[2]/div[5]/div/div/div[2]/div/div[1]/div/div[1]/input')
            enter_oxygensat.send_keys(self.oxygen_sat)

            enter_rr = web.find_element("xpath", '//*[@id="mG61Hd"]/div[2]/div/div[2]/div[6]/div/div/div[2]/div/div[1]/div/div[1]/input')
            enter_rr.send_keys(self.respiratory_rate)

            enter_summary = web.find_element("xpath", '//*[@id="mG61Hd"]/div[2]/div/div[2]/div[7]/div/div/div[2]/div/div[1]/div[2]/textarea')
            enter_summary.send_keys(self.summary)

                #Submit form

            submit = web.find_element("xpath", '//*[@id="mG61Hd"]/div[2]/div/div[3]/div[1]/div[1]/div/span/span')
            submit.click()
            time.sleep(2)
            web.quit()

            # Create a new Word document
            doc = Document()
                    
                    # Add title with patient's name
            doc.add_heading(f'Patient Vital Signs - {self.patient_name}', level=1)
                    
                    # Add vital signs information

            doc.add_paragraph(f"Patient Name: {self.patient_name}")
            doc.add_paragraph(f"Blood pressure (mmHg): {self.bp}")
            doc.add_paragraph(f"Temperature: {self.temperature}")
            doc.add_paragraph(f"Pulse rate: {self.pulse_rate}")
            doc.add_paragraph(f"Oxygen Saturation: {self.oxygen_sat}")
            doc.add_paragraph(f"Respiratory rate: {self.respiratory_rate}")
            doc.add_paragraph(f"Summary of condition: {self.summary}")
                    
            # Save the document
            doc.save(f'{self.patient_name}_vital_signs.docx')

            Clock.schedule_once(self.themainmenu, 1)

    def recall_vital_signs(self, instance):
        self.clear_widgets()
        layout = BoxLayout(orientation='vertical', spacing = 80)
        self.add_widget(layout)

        self.image_widget = Image(source='Your Image Here', 
        allow_stretch=True, 
        keep_ratio=True, 
        size_hint=(None, None)
        )
        layout.add_widget(self.image_widget)
        self.image_widget.pos_hint = {'center_x': 0.5, 'center_y': 0.5}
        self.title_label = Label(text='Patient Name', font_size=64, font_name='Roboto', bold=True, color=(0.12, 0.46, 0.70, 1))
        self.title_label.size_hint = (1, 0.4)
        layout.add_widget(self.title_label)
        self.patient_name_input = TextInput(multiline=False, hint_text='Enter Your Name', size_hint=(0.3, 0.3), pos_hint={'center_x': 0.5, 'center_y': 0.9})
        layout.add_widget(self.patient_name_input)
        
        bottom_layout = GridLayout(cols=2, spacing=1)
        layout.add_widget(bottom_layout)

        left_bottom_layout = BoxLayout(orientation='vertical', spacing=10)
        bottom_layout.add_widget(left_bottom_layout)

        self.back_button = Button(text='Main Menu', size_hint=(0.35, None), pos_hint={'center_x': 0.5, 'center_y': 0.5}, background_color=(0.12, 0.46, 0.70, 1))
        self.back_button.bind(on_press=self.themainmenu)
        left_bottom_layout.add_widget(self.back_button)

        right_bottom_layout = BoxLayout(orientation='vertical', spacing=10)
        bottom_layout.add_widget(right_bottom_layout)

        self.fetch_button = Button(text='Get Vitals', size_hint=(0.35, None), pos_hint={'center_x': 0.5, 'center_y': 0.5}, background_color=(0.12, 0.46, 0.70, 1))
        self.fetch_button.bind(on_press=self.get_patient_name)
        right_bottom_layout.add_widget(self.fetch_button)

        spacer = Widget(size_hint_y=0.4)
        self.add_widget(spacer)


    def get_patient_name(self, instance):
        self.patient_name = self.patient_name_input.text
        self.clear_widgets()
        layout = BoxLayout(orientation='vertical', spacing = 10)
        self.add_widget(layout)  # Add the layout to the screen after adding the label
        sleep(2)
        try:
            doc_path = f'{self.patient_name}_vital_signs.docx'
            os.startfile(doc_path)
        except FileNotFoundError:
            self.clear_widgets()
            label = Label(text='''File not found.
    Redirecting to main menu''', 
                        font_name='Roboto', bold=True, 
                        color=(0, 0, 0, 1), 
                        font_size=30, 
                        pos_hint={'center_x': 0.5, 'center_y': 0.5})
            layout.add_widget(label)
            Clock.schedule_once(self.themainmenu, 3)
        except Exception as e:
            self.clear_widgets()
            label = Label(text='''An error occured.
    Redirecting to main menu''', 
                        font_name='Roboto', bold=True, 
                        color=(0, 0, 0, 1), 
                        font_size=30, 
                        pos_hint={'center_x': 0.5, 'center_y': 0.5})
            layout.add_widget(label)
            Clock.schedule_once(self.themainmenu, 3)
        Clock.schedule_once(self.themainmenu, 3)

    def exit_app(self, instance):
        App.get_running_app().stop()


class VitalSignsApp(App):
    def build(self):
        return MainMenu()

if __name__ == '__main__':
    VitalSignsApp().run()